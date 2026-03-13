import fitz
import base64
import openai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.config import API_KEY, BASE_URL
from docx import Document
from docx.shared import Pt, RGBColor
import re


def insert_markdown_table(doc, markdown_text):
    """简单的解析器：将 Markdown 表格转为 Word 表格"""
    lines = [l.strip() for l in markdown_text.strip().split('\n') if '|' in l]
    if len(lines) < 2:
        doc.add_paragraph(markdown_text)
        return

    # 提取表头和数据（跳过分隔行 ---|---）
    rows = []
    for line in lines:
        if re.match(r'^[|\s:-]+$', line): continue
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells: rows.append(cells)

    if not rows: return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            table.cell(i, j).text = cell_text


def generate_dual_reports(data, img_results, save_path):
    doc = Document()
    # ... (前面的标题、元数据等逻辑保持不变) ...

    # 重点：修改核心结果展示
    doc.add_heading('4. 主要发现与结果 (Key Results)', level=1)
    res = data['key_results']
    doc.add_paragraph(f"• 核心结论：{res['core_conclusion']}")

    # 渲染 Markdown 表格
    doc.add_paragraph("• 关键指标对比：")
    insert_markdown_table(doc, res['data_table'])

    doc.add_paragraph(f"• 图表解读：{res['interpretation']}")

    # 6. 批判性分析 (确保中文术语准确)
    doc.add_heading('6. 个人批判性评价 (Critical Thinking)', level=1)
    crit = data['critical_analysis']
    doc.add_paragraph(f"⭐ 创新评估：{crit['innovation']}")
    doc.add_paragraph(f"🎯 对我课题 (SNN Co-design) 的启发：{crit['relevance_to_snn_codesign']}")

    doc.save(save_path / "1_深度总结报告.docx")

client = openai.OpenAI(api_key=API_KEY, base_url=BASE_URL)


def extract_and_analyze_images(pdf_path, output_folder):
    """提取 PDF 中的大尺寸图片并利用 GPT-4o 进行视觉分析"""
    img_dir = output_folder / "Figures"
    img_dir.mkdir(exist_ok=True)

    doc = fitz.open(pdf_path)
    image_analyses = []

    for page_index in range(len(doc)):
        images = doc[page_index].get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]

            # 过滤杂讯：忽略小于 20KB 的图片（通常是图标、页码或Logo）
            if len(image_bytes) < 20000: continue

            img_name = f"page{page_index + 1}_img{img_index + 1}.png"
            img_path = img_dir / img_name
            with open(img_path, "wb") as f:
                f.write(image_bytes)

            # 使用 GPT-4o 进行视觉分析
            try:
                base64_img = base64.b64encode(image_bytes).decode('utf-8')
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "user", "content": [
                        {"type": "text",
                         "text": "分析这张学术图表。1.它是哪类图？2.关键实验结果或架构细节是什么？3.在SNN协同设计中有什么意义？"},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_img}"}}
                    ]}],
                    max_tokens=400
                )
                image_analyses.append({
                    "path": img_name,
                    "desc": response.choices[0].message.content
                })
            except Exception as e:
                print(f"⚠️ 图片分析失败: {img_name}, 错误: {e}")
                continue
    return image_analyses


def generate_dual_reports(analysis_json, img_results, save_path):
    """根据深度 JSON 数据生成两份 Word 报告"""

    # --- 1. 深度总结报告渲染 ---
    doc_sum = Document()

    # 样式设置：标题
    title = doc_sum.add_heading(analysis_json['bibliographic_info']['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一句话总结
    p_high = doc_sum.add_paragraph()
    run = p_high.add_run(f"💡 一句话总结：{analysis_json['one_sentence_summary']}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)  # 深蓝色

    # 1. 基础元数据
    doc_sum.add_heading('1. 基础元数据 (Bibliographic Info)', level=1)
    meta = analysis_json['bibliographic_info']
    doc_sum.add_paragraph(f"• 作者与年份：{meta['authors_year']}")
    doc_sum.add_paragraph(f"• 来源：{meta['source']}")
    doc_sum.add_paragraph(f"• 关键词：{', '.join(meta['keywords'])}")

    # 2. 研究背景与动机
    doc_sum.add_heading('2. 研究背景与动机', level=1)
    bm = analysis_json['background_motivation']
    doc_sum.add_paragraph(f"• 研究缺口：{bm['research_gap']}")
    doc_sum.add_paragraph(f"• 核心问题：{bm['research_question']}")
    doc_sum.add_paragraph(f"• 重要性：{bm['importance']}")

    # 3. 核心方法论
    doc_sum.add_heading('3. 核心方法论 (Methodology)', level=1)
    meth = analysis_json['methodology']
    doc_sum.add_paragraph(f"• 技术创新：{meth['key_tech']}")
    doc_sum.add_paragraph(f"• 数据集/基准：{meth['data_source']} (基准: {meth['baselines']})")

    # 4. 主要发现与结果
    doc_sum.add_heading('4. 主要发现与结果 (Key Results)', level=1)
    res = analysis_json['key_results']
    doc_sum.add_paragraph(f"• 核心结论：{res['core_conclusion']}")
    doc_sum.add_paragraph(f"• 关键指标：{res['key_metrics']}")
    doc_sum.add_paragraph(f"• 图表解读：{res['visual_interpretation']}")

    # 5. 讨论与局限性
    doc_sum.add_heading('5. 讨论与局限性', level=1)
    dis = analysis_json['discussion_limitations']
    doc_sum.add_paragraph(f"• 局限性：{dis['limitations']}")
    doc_sum.add_paragraph(f"• 未来展望：{dis['future_work']}")

    # 6. 个人批判性评价 (Critical Thinking)
    doc_sum.add_heading('6. 个人批判性评价', level=1)
    crit = analysis_json['critical_analysis']
    doc_sum.add_paragraph(f"⭐ 创新评估：{crit['innovation_eval']}")
    doc_sum.add_paragraph(f"🔍 可靠性评估：{crit['reliability_eval']}")
    p_insp = doc_sum.add_paragraph(f"🎯 对我课题的启发：{crit['relevance_to_user']}")
    p_insp.paragraph_format.space_before = Pt(10)

    # 7. 金句摘录
    doc_sum.add_heading('7. 金句摘录 (Key Quotes)', level=1)
    for quote in analysis_json['key_quotes']:
        q_p = doc_sum.add_paragraph(quote, style='Quote')

    doc_sum.save(save_path / "1_深度总结报告.docx")

    # --- 2. 图片解析报告渲染 ---
    doc_img = Document()
    doc_img.add_heading(f"图表详细解析: {analysis_json['bibliographic_info']['authors_year']}", 0)

    if img_results:
        for img in img_results:
            doc_img.add_heading(f"📷 图片: {img['path']}", level=2)
            doc_img.add_paragraph(img['desc'])
            doc_img.add_paragraph("-" * 20)
    else:
        doc_img.add_paragraph("在文中未检测到有效的大尺寸学术图表。")

    doc_img.save(save_path / "2_图表解析报告.docx")